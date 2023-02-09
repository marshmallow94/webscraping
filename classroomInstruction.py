import requests
from bs4 import BeautifulSoup
import docx
from docx.shared import Cm
from urllib3.exceptions import InsecureRequestWarning
from urllib3 import disable_warnings

headings = ['TURNING ON THE SYSTEM', 'USING THE', 'DTEN Quickstart guide']
laptop = ['CONNECTING A LAPTOP:', 'To display a laptop' ]
cablecast = 'RECEIVING CABLECAST'

obtained = False
special = False
footer_text = ""

#disabling warnings 
disable_warnings(InsecureRequestWarning)

#create document and fomatt it
document = docx.Document()
sections = document.sections
for s in sections:
    s.top_margin = Cm(1)
    s.bottom_margin = Cm(1)
    s.left_margin = Cm(1)
    s.right_margin = Cm(1)

section = document.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]

style = document.styles['Normal']
font = style.font
font.name = 'Times'
font.size = docx.shared.Pt(11)

#get data from url
response = requests.get("https://classrooms.sfsu.edu/rooms", verify=False)

#start scraping
soup = BeautifulSoup(response.text, 'html.parser')
cards = soup.find_all("div", {"class": "panel panel-default room-card"})

#get the room number from each room cards 
for num in cards:
    room = num.find('h3').text
    title = document.add_heading('ROOM ' + room.strip())
    title.add_run().bold = True
    title.style.font.name = 'Times'
    title.style.font.size = docx.shared.Pt(14)

    link = num.find('div', {"class": "col-sm-3 tutorial-info"}).find('a')
    
    #check if the room has tutorial link 
    if(link != None):
        tutorial = link.get('href')
        res = requests.get(tutorial, verify=False)
        sp = BeautifulSoup(res.text, 'html.parser')

        if not obtained:
            footer_text = sp.find('div', {'class': "well"}).text
            footer_para.text = '\t\t' + footer_text
            obtained = True

        body = sp.find('div', {'class': "tutorial-body"})
        
        if('MH' in room) or ('BH' in room):
            print(body.text)
        
        skip = False
        body_text = body.text.replace(u'\xa0', u' ').split('\n')
        count = 0
        

        for text in body_text:
            clean_text = text.strip()

            if  [ h for h in headings if ( h in clean_text )]:
                heading = document.add_heading(clean_text)
                heading_style = heading.style
                heading_style.font.name = "Times"
                skip = False
            elif [ l for l in laptop if ( l in clean_text )] :
                heading = document.add_heading(clean_text)
                heading_style = heading.style
                heading_style.font.name = "Times"
                special = True
                skip = False
            elif clean_text == cablecast:
                skip = True
                continue
            elif text == '\r' or text == '' or skip:
                continue
            elif special and 'Mirroring' in clean_text:
                a = clean_text.replace('Mirroring',u'\n\nMirroring')
                b = a.replace('Mac:', u':\nMac:')
                c = b.replace('Windows:', u':\nWindows:')
                document.add_paragraph(c)
                special = False
            else:
                document.add_paragraph(clean_text)

    else:
        document.add_paragraph('[No tutorial on website]')
    document.add_page_break()


# save the document
document.save('/Users/PATH'+ 'classroomTutorial.docx')
